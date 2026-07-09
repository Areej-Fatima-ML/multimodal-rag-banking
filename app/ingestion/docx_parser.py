from docx import Document
import io
from PIL import Image


def extract_text_from_docx(file_path: str) -> list:
    """
    Extract all paragraphs text from DOCX file
    Returns list of text dicts with content and metadata
    """
    texts = []
    doc = Document(file_path)

    for para_num, paragraph in enumerate(doc.paragraphs):
        # Only add non-empty paragraphs
        if paragraph.text.strip():
            texts.append({
                "content": paragraph.text,
                "paragraph": para_num + 1,
                "type": "text",
                "source": file_path
            })

    return texts


def extract_tables_from_docx(file_path: str) -> list:
    """
    Extract all tables from DOCX file
    Converts table data to readable text format
    Returns list of table dicts with content and metadata
    """
    tables = []
    doc = Document(file_path)

    for table_num, table in enumerate(doc.tables):
        table_text = ""

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_text += " | ".join(row_data) + "\n"

        # Only add non-empty tables
        if table_text.strip():
            tables.append({
                "content": table_text,
                "table_num": table_num + 1,
                "type": "table",
                "source": file_path
            })

    return tables


def extract_images_from_docx(file_path: str) -> list:
    """
    Extract all embedded images from DOCX file
    Returns list of image dicts with PIL image and metadata
    """
    images = []
    doc = Document(file_path)

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            image = Image.open(io.BytesIO(image_data))

            # Skip very small images
            if image.width >= 10 and image.height >= 10:
                images.append({
                    "image": image,
                    "type": "image",
                    "source": file_path
                })

    return images


def parse_docx(file_path: str) -> dict:
    """
    Parse complete DOCX file
    Extracts text, tables and images
    Returns dict with all extracted content and metadata
    """
    texts = extract_text_from_docx(file_path)
    tables = extract_tables_from_docx(file_path)
    images = extract_images_from_docx(file_path)

    return {
        "file_path": file_path,
        "texts": texts,
        "tables": tables,
        "images": images,
        "total_texts": len(texts),
        "total_tables": len(tables),
        "total_images": len(images)
    }